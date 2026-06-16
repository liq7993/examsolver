"""Generate editable Word notes with native OMML equations."""

from __future__ import annotations

import json
import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from examsolver.contracts import NoteEntry

_COMMAND_TEXT = {
    "alpha": "α",
    "beta": "β",
    "cdot": "·",
    "cos": "cos",
    "gamma": "γ",
    "ln": "ln",
    "pi": "π",
    "sin": "sin",
    "tan": "tan",
    "theta": "θ",
    "times": "×",
}


def export_note_to_docx(note: NoteEntry) -> bytes:
    """Render one NoteEntry into an in-memory docx file."""

    document = Document()
    document.styles["Normal"].font.name = "Microsoft YaHei"
    document.styles["Normal"].font.size = Pt(11)

    title = document.add_heading(note.title or "单题笔记", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"学科：{note.subject or 'unknown'}    题型：{note.question_type}")

    document.add_heading("题目", level=1)
    document.add_paragraph(note.question_latex)

    explanation = note.student_explanation
    if explanation is not None:
        document.add_heading("思路", level=1)
        document.add_paragraph(explanation.intuition or explanation.summary)

    document.add_heading("步骤", level=1)
    for step in note.steps:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(step.description)
        if step.formula_latex:
            _add_equation(document.add_paragraph(), step.formula_latex)

    document.add_heading("答案", level=1)
    answer_paragraph = document.add_paragraph()
    if isinstance(note.answer, str):
        _add_equation(answer_paragraph, note.answer)
    elif note.answer is None:
        answer_paragraph.add_run("暂无答案")
    else:
        answer_paragraph.add_run(json.dumps(note.answer, ensure_ascii=False, indent=2))

    document.add_heading("易错点", level=1)
    mistakes = list(note.common_mistakes)
    if explanation is not None and explanation.common_mistake:
        if explanation.common_mistake not in mistakes:
            mistakes.append(explanation.common_mistake)
    if mistakes:
        for mistake in mistakes:
            document.add_paragraph(mistake, style="List Bullet")
    else:
        document.add_paragraph("暂无易错点。")

    if note.related_formulas:
        document.add_heading("公式速查", level=1)
        for formula in note.related_formulas:
            document.add_paragraph(formula.title, style="Heading 2")
            _add_equation(document.add_paragraph(), formula.formula_latex)
            if formula.explanation:
                document.add_paragraph(formula.explanation)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_equation(paragraph: Any, latex: str) -> None:
    source = _strip_math_delimiters(latex)
    math = OxmlElement("m:oMath")
    try:
        for element in _LatexParser(source).parse():
            math.append(element)
    except ValueError:
        math.append(_math_run(source))
    paragraph._p.append(math)


def _strip_math_delimiters(value: str) -> str:
    source = value.strip()
    for left, right in (("$$", "$$"), ("$", "$"), ("\\[", "\\]"), ("\\(", "\\)")):
        if source.startswith(left) and source.endswith(right):
            return source[len(left) : -len(right)].strip()
    return source


class _LatexParser:
    def __init__(self, source: str) -> None:
        self.source = source
        self.position = 0

    def parse(self, stop: str | None = None) -> list[Any]:
        elements: list[Any] = []
        while self.position < len(self.source):
            if stop is not None and self.source[self.position] == stop:
                self.position += 1
                return elements
            elements.extend(self._parse_atom())
        if stop is not None:
            raise ValueError("unclosed latex group")
        return elements

    def _parse_atom(self) -> list[Any]:
        char = self.source[self.position]
        if char == "{":
            self.position += 1
            base = self.parse("}")
        elif char == "\\":
            base = self._parse_command()
        else:
            self.position += 1
            base = [_math_run(char)]

        while self.position < len(self.source) and self.source[self.position] in "^_":
            operator = self.source[self.position]
            self.position += 1
            script = self._parse_script_value()
            base = [_script_element(base, script, superscript=operator == "^")]
        return base

    def _parse_command(self) -> list[Any]:
        self.position += 1
        match = re.match(r"[A-Za-z]+", self.source[self.position :])
        if match is None:
            if self.position >= len(self.source):
                return []
            char = self.source[self.position]
            self.position += 1
            return [_math_run(char)]

        command = match.group(0)
        self.position += len(command)
        if command == "frac":
            return [_fraction_element(self._parse_required_group(), self._parse_required_group())]
        if command == "sqrt":
            return [_radical_element(self._parse_required_group())]
        if command == "text":
            return [_math_run(_elements_text(self._parse_required_group()))]
        if command in {"left", "right"}:
            return []
        return [_math_run(_COMMAND_TEXT.get(command, command))]

    def _parse_required_group(self) -> list[Any]:
        self._skip_spaces()
        if self.position >= len(self.source) or self.source[self.position] != "{":
            raise ValueError("latex command requires a group")
        self.position += 1
        return self.parse("}")

    def _parse_script_value(self) -> list[Any]:
        self._skip_spaces()
        if self.position >= len(self.source):
            raise ValueError("missing script value")
        if self.source[self.position] == "{":
            self.position += 1
            return self.parse("}")
        return self._parse_atom()

    def _skip_spaces(self) -> None:
        while self.position < len(self.source) and self.source[self.position].isspace():
            self.position += 1


def _math_run(text: str) -> Any:
    run = OxmlElement("m:r")
    value = OxmlElement("m:t")
    value.set(qn("xml:space"), "preserve")
    value.text = text
    run.append(value)
    return run


def _fraction_element(numerator: list[Any], denominator: list[Any]) -> Any:
    fraction = OxmlElement("m:f")
    num = OxmlElement("m:num")
    den = OxmlElement("m:den")
    for item in numerator:
        num.append(item)
    for item in denominator:
        den.append(item)
    fraction.extend((num, den))
    return fraction


def _radical_element(content: list[Any]) -> Any:
    radical = OxmlElement("m:rad")
    radical.append(OxmlElement("m:deg"))
    body = OxmlElement("m:e")
    for item in content:
        body.append(item)
    radical.append(body)
    return radical


def _script_element(base: list[Any], script: list[Any], *, superscript: bool) -> Any:
    element = OxmlElement("m:sSup" if superscript else "m:sSub")
    body = OxmlElement("m:e")
    script_body = OxmlElement("m:sup" if superscript else "m:sub")
    for item in base:
        body.append(item)
    for item in script:
        script_body.append(item)
    element.extend((body, script_body))
    return element


def _elements_text(elements: list[Any]) -> str:
    return "".join("".join(element.itertext()) for element in elements)
